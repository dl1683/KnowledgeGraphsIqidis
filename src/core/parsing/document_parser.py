"""
Document parsing for PDF and DOCX files.
"""
import hashlib
from pathlib import Path
from typing import Optional, Dict, Any
from dataclasses import dataclass


@dataclass
class ParsedDocument:
    """Result of parsing a document."""
    filepath: str
    filename: str
    file_hash: str
    text: str
    metadata: Dict[str, Any]
    page_count: int = 0


class DocumentParser:
    """Parse PDF and DOCX documents to extract text."""

    def __init__(self):
        self._check_dependencies()

    def _check_dependencies(self):
        """Check for required parsing libraries."""
        self.has_pymupdf = False
        self.has_docx = False

        try:
            import fitz  # PyMuPDF
            self.has_pymupdf = True
        except ImportError:
            pass

        try:
            import docx
            self.has_docx = True
        except ImportError:
            pass

    def parse(self, filepath: str) -> Optional[ParsedDocument]:
        """Parse a document and return extracted text."""
        path = Path(filepath)

        if not path.exists():
            print(f"File not found: {filepath}")
            return None

        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self._parse_pdf(path)
        elif suffix in [".docx", ".doc"]:
            return self._parse_docx(path)
        elif suffix == ".txt":
            return self._parse_txt(path)
        else:
            print(f"Unsupported file type: {suffix}")
            return None

    def _parse_pdf(self, path: Path) -> Optional[ParsedDocument]:
        """Parse a PDF file."""
        if not self.has_pymupdf:
            print("PyMuPDF not installed. Install with: pip install PyMuPDF")
            return None

        import fitz

        try:
            doc = fitz.open(str(path))
            text_parts = []

            for page_num, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")

            full_text = "\n\n".join(text_parts)

            metadata = {
                "page_count": len(doc),
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", ""),
                "creator": doc.metadata.get("creator", ""),
            }

            file_hash = self._compute_hash(path)
            doc.close()

            return ParsedDocument(
                filepath=str(path),
                filename=path.name,
                file_hash=file_hash,
                text=full_text,
                metadata=metadata,
                page_count=metadata["page_count"]
            )

        except Exception as e:
            print(f"Error parsing PDF {path}: {e}")
            return None

    def _parse_docx(self, path: Path) -> Optional[ParsedDocument]:
        """Parse a DOCX file."""
        if not self.has_docx:
            print("python-docx not installed. Install with: pip install python-docx")
            return None

        from docx import Document as DocxDocument

        try:
            doc = DocxDocument(str(path))
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            full_text = "\n\n".join(paragraphs)

            metadata = {
                "author": doc.core_properties.author or "",
                "title": doc.core_properties.title or "",
                "subject": doc.core_properties.subject or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
            }

            file_hash = self._compute_hash(path)

            return ParsedDocument(
                filepath=str(path),
                filename=path.name,
                file_hash=file_hash,
                text=full_text,
                metadata=metadata,
                page_count=0  # DOCX doesn't have a clear page concept
            )

        except Exception as e:
            print(f"Error parsing DOCX {path}: {e}")
            return None

    def _parse_txt(self, path: Path) -> Optional[ParsedDocument]:
        """Parse a plain text file."""
        try:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()

            file_hash = self._compute_hash(path)

            return ParsedDocument(
                filepath=str(path),
                filename=path.name,
                file_hash=file_hash,
                text=text,
                metadata={},
                page_count=0
            )

        except Exception as e:
            print(f"Error parsing TXT {path}: {e}")
            return None

    def _compute_hash(self, path: Path) -> str:
        """Compute SHA256 hash of file."""
        sha256_hash = hashlib.sha256()
        with open(path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()

    def get_supported_extensions(self) -> list:
        """Get list of supported file extensions."""
        extensions = [".txt"]
        if self.has_pymupdf:
            extensions.append(".pdf")
        if self.has_docx:
            extensions.extend([".docx", ".doc"])
        return extensions
